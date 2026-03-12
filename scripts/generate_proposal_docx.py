"""
generate_proposal_docx.py

Generates a dissertation-quality research proposal in Bahasa Indonesia as a .docx file.
Output: docs/proposal_penelitian.docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def set_style(doc):
    """Set document-wide styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Cm(1.25)

    for heading in ["Heading 1", "Heading 2", "Heading 3"]:
        if heading in doc.styles:
            h = doc.styles[heading]
            h.font.name = "Times New Roman"
            h.font.size = Pt(14 if "1" in heading else 12)
            h.font.bold = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_paragraph(doc, text):
    """Add a properly formatted paragraph."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_ref_list_item(doc, text, hanging=True):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25) if hanging else Cm(0)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def main():
    doc = Document()
    set_style(doc)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ========== HALAMAN JUDUL ==========
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PENGEMBANGAN SISTEM REKOMENDASI KURIKULUM\nREKAYASA PERANGKAT LUNAK JENJANG SMK\nBERBASIS ANALISIS KEBUTUHAN PASAR TENAGA KERJA\nYANG FUTURE-AWARE")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.all_caps = False

    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Proposal Penelitian Disertasi")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(12)

    doc.add_page_break()

    # ========== BAB I PENDAHULUAN ==========
    add_heading(doc, "BAB I", level=1)
    add_heading(doc, "PENDAHULUAN", level=1)

    # 1.1 Latar Belakang Masalah
    add_heading(doc, "1.1 Latar Belakang Masalah", level=2)

    add_paragraph(doc,
        "Pendidikan kejuruan memegang peran strategis dalam mempersiapkan tenaga kerja terampil "
        "yang sesuai dengan kebutuhan industri. Dalam konteks Indonesia, Sekolah Menengah Kejuruan (SMK) "
        "merupakan institusi utama yang menghasilkan lulusan siap kerja di berbagai sektor. Namun, "
        "data Badan Pusat Statistik (BPS, 2024) menunjukkan bahwa tingkat pengangguran terbuka (TPT) "
        "lulusan SMK mencapai 9,01% per Agustus 2024—angka tertinggi dibandingkan dengan jenjang pendidikan "
        "lainnya seperti SMA (7,05%), Diploma (4,83%), dan perguruan tinggi (5,25%).")

    add_paragraph(doc,
        "Fenomena tingginya pengangguran lulusan SMK tidak terlepas dari kesenjangan antara kompetensi "
        "yang diajarkan di sekolah dengan kebutuhan aktual dunia kerja. Saputri dan Sulistyawati (2024) "
        "dalam studi terhadap 101.748 lulusan SMK di Indonesia menemukan bahwa 61,58% mengalami horizontal "
        "mismatch—ketrampilan yang dimiliki tidak selaras dengan persyaratan pekerjaan. Sebanyak 13,58% "
        "mengalami overeducation, dan 10,13% mengalami real mismatch (kombinasi keduanya). Temuan ini "
        "menunjukkan bahwa desain kurikulum yang tidak selaras dengan pasar tenaga kerja menjadi akar "
        "permasalahan utama.")

    add_paragraph(doc,
        "World Economic Forum (2025) dalam Future of Jobs Report 2025 menegaskan bahwa kesenjangan "
        "keterampilan (skill gap) merupakan hambatan utama transformasi bisnis—disebutkan oleh 63% "
        "employer untuk periode 2025–2030. Keterampilan yang paling cepat berkembang meliputi AI dan "
        "big data, jaringan dan cybersecurity, serta literasi teknologi. Laporan tersebut juga "
        "mengidentifikasi pemikiran kreatif, ketahanan, fleksibilitas, dan kepemimpinan sebagai "
        "keterampilan penting yang semakin dibutuhkan. Implikasinya, kurikulum pendidikan kejuruan "
        "harus mampu beradaptasi dengan dinamika pasar kerja yang berubah cepat.")

    add_paragraph(doc,
        "Indonesia telah mengadopsi Kurikulum Merdeka melalui Peraturan Menteri Pendidikan, Kebudayaan, "
        "Riset, dan Teknologi Nomor 56 Tahun 2022, yang memberikan otonomi kepada satuan pendidikan "
        "untuk mengembangkan Kurikulum Operasional Satuan Pendidikan (KOSP) sesuai konteks dan kebutuhan "
        "lokal. Prinsip pengembangan KOSP mencakup kontekstualitas—yaitu menjunjukkan kekhasan satuan "
        "pendidikan sesuai karakteristik dunia kerja dan industri (khusus SMK)—serta akuntabilitas "
        "berbasis data yang aktual (Permendikbudristek, 2022). Namun, dalam praktiknya, pengembangan "
        "KOSP masih banyak mengandalkan penilaian subjektif dan pengalaman guru, tanpa dukungan sistematis "
        "dari analisis data pasar tenaga kerja.")

    add_paragraph(doc,
        "Teori human capital yang dikembangkan Becker (1964) dan Schultz (1961) menempatkan pendidikan "
        "sebagai investasi yang menghasilkan return berupa peningkatan produktivitas dan pendapatan. "
        "Investasi dalam keterampilan yang selaras dengan permintaan pasar tenaga kerja akan memaksimalkan "
        "return tersebut. Oleh karena itu, kurikulum yang berbasis data analisis kebutuhan pasar tenaga "
        "kerja tidak hanya relevan secara pedagogis, tetapi juga ekonomis.")

    add_paragraph(doc,
        "Perkembangan kecerdasan buatan (AI), khususnya Natural Language Processing (NLP) dan Large "
        "Language Models (LLM), membuka peluang baru untuk menganalisis data pasar tenaga kerja secara "
        "otomatis. Zhang et al. (2022) memperkenalkan SkillSpan—dataset dan model ekstraksi keterampilan "
        "dari lowongan kerja berbasis BERT. Senger et al. (2024) melakukan survei komprehensif mengenai "
        "metode deep learning untuk ekstraksi dan klasifikasi keterampilan dari job postings. Penelitian "
        "terbaru juga menunjukkan aplikasi LLM dalam desain kurikulum (ACL 2025) dan sistem rekomendasi "
        "pendidikan (ArXiv, 2024). Pendekatan hybrid yang menggabungkan model NER berbasis BERT dengan "
        "LLM dapat mengatasi keterbatasan masing-masing: BERT unggul untuk ekstraksi terstruktur di level "
        "kalimat, sementara LLM menyediakan pemahaman kontekstual pada teks panjang.")

    add_paragraph(doc,
        "Berdasarkan uraian di atas, penelitian ini mengusulkan pengembangan sistem rekomendasi kurikulum "
        "Rekayasa Perangkat Lunak (RPL) jenjang SMK yang future-aware—yakni memanfaatkan ekstraksi "
        "keterampilan dari data lowongan kerja dengan teknik NLP hybrid, analisis tren temporal yang "
        "dikontrol secara statistik (FDR), dan pemetaan ke domain masa depan berdasarkan laporan WEF dan "
        "O*NET. Sistem ini diharapkan dapat memberikan rekomendasi berbasis data yang actionable bagi "
        "sekolah dalam mengembangkan KOSP sesuai prinsip Kurikulum Merdeka.")

    doc.add_paragraph()

    # 1.2 Rumusan Masalah
    add_heading(doc, "1.2 Rumusan Masalah", level=2)

    add_paragraph(doc,
        "Berdasarkan latar belakang di atas, rumusan masalah penelitian ini adalah:")

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("1. ").bold = True
    p.add_run(
        "Bagaimana mengembangkan sistem rekomendasi kurikulum RPL SMK yang future-aware dengan "
        "memanfaatkan ekstraksi keterampilan berbasis NLP hybrid (BERT + LLM) dari data pasar tenaga kerja?")

    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.add_run("2. ").bold = True
    p2.add_run(
        "Seberapa efektif sistem tersebut dalam mengidentifikasi keterampilan yang relevan dan "
        "selaras dengan kebutuhan pasar tenaga kerja (dihitung melalui precision ekstraksi, "
        "akurasi pemetaan domain, dan metrik evaluasi saintifik)?")

    p3 = doc.add_paragraph()
    p3.paragraph_format.first_line_indent = Cm(0)
    p3.add_run("3. ").bold = True
    p3.add_run(
        "Bagaimana persepsi stakeholders (guru, pengembang kurikulum, kepala sekolah) terhadap "
        "kegunaan dan penerapan sistem rekomendasi kurikulum berbasis data tersebut?")

    doc.add_paragraph()

    # 1.3 Tujuan Penelitian
    add_heading(doc, "1.3 Tujuan Penelitian", level=2)

    add_paragraph(doc,
        "Tujuan penelitian ini adalah:")

    for i, tujuan in enumerate([
        "Mengembangkan dan memvalidasi sistem rekomendasi kurikulum RPL SMK yang memanfaatkan ekstraksi "
        "keterampilan hybrid (JobBERT + LLM) dari data lowongan kerja, analisis tren FDR-controlled, "
        "dan pemetaan domain masa depan berbasis WEF/O*NET.",
        "Mengukur efektivitas sistem melalui metrik saintifik: precision ekstraksi (Wilson CI, binomial test), "
        "akurasi pemetaan domain (Top-1, Top-3, MRR), serta validasi Bloom taxonomy dan kompetensi.",
        "Mengevaluasi penerimaan stakeholders (guru, pengembang kurikulum) terhadap kegunaan sistem "
        "melalui expert review dan wawancara.",
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(f"{i}. ").bold = True
        p.add_run(tujuan)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)

    doc.add_paragraph()

    # 1.4 Manfaat Penelitian
    add_heading(doc, "1.4 Manfaat Penelitian", level=2)

    add_paragraph(doc,
        "Manfaat penelitian dapat dikategorikan sebagai berikut:")

    add_paragraph(doc,
        "Manfaat teoretis: Penelitian ini memperluas teori human capital dengan menunjukkan bagaimana "
        "AI dapat menjadi mediator antara sinyal pasar tenaga kerja dan desain kurikulum. Temuan "
        "mengenai efektivitas pendekatan hybrid NLP untuk ekstraksi keterampilan juga berkontribusi "
        "pada literatur computational job market analysis (Senger et al., 2024).")

    add_paragraph(doc,
        "Manfaat praktis: Sistem yang dikembangkan memberikan alat actionable bagi SMK—khususnya "
        "program keahlian RPL—untuk mengembangkan KOSP berbasis data. Rekomendasi prioritas kurikulum "
        "(demand, trend, future-weight) dapat langsung digunakan dalam proses pengambilan keputusan "
        "kurikuler.")

    add_paragraph(doc,
        "Manfaat kebijakan: Hasil penelitian dapat menginformasikan Kemendikbudristek mengenai "
        "potensi data-driven KOSP development, serta perlunya dukungan infrastruktur dan kapasitas "
        "untuk adopsi pendekatan serupa di skala nasional.")

    doc.add_page_break()

    # ========== BAB II KAJIAN PUSTAKA ==========
    add_heading(doc, "BAB II", level=1)
    add_heading(doc, "KAJIAN PUSTAKA", level=1)

    # 2.1 Teori Human Capital
    add_heading(doc, "2.1 Teori Human Capital", level=2)

    add_paragraph(doc,
        "Teori human capital yang dikembangkan Becker (1964) dan Schultz (1961) menempatkan keterampilan "
        "dan pengetahuan sebagai bentuk modal yang dihasilkan dari investasi sebelumnya. Schultz (1961) "
        "menyatakan bahwa kegiatan seperti pendidikan, pelatihan kerja, dan perawatan kesehatan "
        "meningkatkan kemampuan fisik dan mental individu, sehingga meningkatkan prospek pendapatan riil. "
        "Becker (1964) mendefinisikan investasi human capital sebagai aktivitas yang mempengaruhi "
        "pendapatan riil di masa depan melalui penanaman sumber daya pada manusia.")

    add_paragraph(doc,
        "Implikasi teori ini bagi pendidikan kejuruan adalah bahwa kurikulum yang selaras dengan "
        "kebutuhan pasar tenaga kerja akan memaksimalkan return investasi pendidikan. Mismatch antara "
        "keterampilan yang diajarkan dan yang dibutuhkan industri mengakibatkan inefisiensi—lulusan "
        "tidak dapat memanfaatkan pendidikan yang telah diinvestasikan. Oleh karena itu, desain "
        "kurikulum berbasis data pasar tenaga kerja bersifat fundamental dari perspektif human capital.")

    # 2.2 Kurikulum Merdeka dan KOSP
    add_heading(doc, "2.2 Kurikulum Merdeka dan Kurikulum Operasional Satuan Pendidikan", level=2)

    add_paragraph(doc,
        "Kurikulum Merdeka diatur melalui Peraturan Menteri Pendidikan, Kebudayaan, Riset, dan "
        "Teknologi Nomor 56 Tahun 2022 (Permendikbudristek, 2022). Kurikulum ini menekankan pembelajaran "
        "berpusat pada peserta didik, kontekstualitas, dan fleksibilitas. Untuk SMK, Kurikulum "
        "Operasional Satuan Pendidikan (KOSP) menjabarkan kurikulum inti ke dalam bentuk konsentrasi "
        "serta potensi internal sekolah dan dunia kerja (Kemendikbudristek, 2022).")

    add_paragraph(doc,
        "Prinsip pengembangan KOSP mencakup: (1) berpusat pada peserta didik; (2) kontekstual—"
        "menunjukkan kekhasan satuan pendidikan sesuai karakteristik dunia kerja dan industri; "
        "(3) esensial—ringkas dan mudah dipahami; (4) akuntabel—berbasis data dan aktual; (5) "
        "melibatkan pemangku kepentingan termasuk industri (Permendikbudristek, 2022). Prinsip "
        "akuntabel dan kontekstual secara eksplisit menuntut KOSP didasarkan pada data aktual "
        "pasar tenaga kerja, namun mekanisme operasional untuk memperoleh dan menganalisis data "
        "tersebut belum tersedia secara sistematis.")

    # 2.3 Kesenjangan Keterampilan dan Pengangguran SMK di Indonesia
    add_heading(doc, "2.3 Kesenjangan Keterampilan dan Pengangguran SMK di Indonesia", level=2)

    add_paragraph(doc,
        "BPS (2024) melaporkan bahwa tingkat pengangguran terbuka lulusan SMK per Agustus 2024 "
        "mencapai 9,01%—tertinggi dibandingkan jenjang lain. Saputri dan Sulistyawati (2024) "
        "menganalisis 101.748 lulusan SMK dan menemukan 61,58% horizontal mismatch, 13,58% "
        "overeducation, dan 10,13% real mismatch. Studi oleh Newhouse dan Suryadarma (2011) "
        "memperlihatkan bahwa lulusan SMK memperoleh pekerjaan lebih cepat dibanding lulusan SMA "
        "namun menghadapi tantangan kesesuaian keterampilan.")

    add_paragraph(doc,
        "Faktor penyebab mismatch meliputi kurikulum yang kurang fleksibel, keterlibatan industri "
        "yang terbatas dalam desain program, disparitas kualitas antardaerah, serta kesenjangan "
        "antara materi yang diajarkan dan kebutuhan employer (Saputri & Sulistyawati, 2024).")

    # 2.4 AI-Assisted Curriculum Design
    add_heading(doc, "2.4 AI-Assisted Curriculum Design", level=2)

    add_paragraph(doc,
        "Penelitian terbaru menunjukkan aplikasi AI dalam desain kurikulum. Studi pada SMK "
        "menemukan bahwa AI-assisted personalized curriculum menghasilkan peningkatan rata-rata "
         "15% pada keterampilan praktis siswa dalam enam bulan (Jurnal Online Informatika, 2024). "
        "AI dan LLM memungkinkan institusi menganalisis tren pasar kerja, mengadaptasi materi "
        "pembelajaran secara real-time, dan memberikan rekomendasi penyesuaian kurikulum "
        "(ISJR, 2024).")

    add_paragraph(doc,
        "MDPI (2024) menerbitkan artikel mengenai AI dalam desain kurikulum perguruan tinggi "
        "dengan pendekatan data-driven. Systematic review mencakup tiga domain: framework "
        "kurikulum adaptif dan etis, literasi AI guru, serta kesiapan institusi dan kebijakan.")

    # 2.5 NLP untuk Ekstraksi Keterampilan dari Job Postings
    add_heading(doc, "2.5 NLP untuk Ekstraksi Keterampilan dari Job Postings", level=2)

    add_paragraph(doc,
        "Zhang et al. (2022) memperkenalkan SkillSpan—dataset dan metode ekstraksi hard dan soft "
        "skill dari lowongan kerja bahasa Inggris. Dataset terdiri dari 14.5K kalimat dengan "
        "lebih dari 12.5K span teranotasi. Penulis mengetes BERT, domain-adapted continuous "
        "pre-training, dan multi-task learning; domain-adapted model secara signifikan mengungguli "
        "model non-adapted.")

    add_paragraph(doc,
        "Senger et al. (2024) menyajikan survei pertama yang berfokus pada NLP untuk ekstraksi "
        "dan klasifikasi keterampilan dari job postings, meliputi metodologi deep learning, "
        "dataset, dan terminologi standar. Survei ini mengidentifikasi pendekatan LLM dengan "
        "synthetic training data sebagai tren emerging.")

    add_paragraph(doc,
        "Pendekatan hybrid yang menggabungkan model NER kecil dengan LLM menggunakan uncertainty "
        "estimation telah dieksplorasi (LinkNER, 2024) untuk mengatasi keterbatasan masing-masing. "
        "Reimers dan Gurevych (2019) mengembangkan Sentence-BERT untuk embedding semantik efisien "
        "yang digunakan dalam pemetaan keterampilan ke domain.")

    # 2.6 Large Language Models dalam Pendidikan
    add_heading(doc, "2.6 Large Language Models dalam Pendidikan", level=2)

    add_paragraph(doc,
        "LLM semakin banyak diterapkan dalam kurikulum dan sistem rekomendasi pendidikan. Studi pada pendidikan tinggi "
        "Finlandia menunjukkan LLM-assisted curriculum writing mengurangi beban kognitif pendidik "
        "sambil mempertahankan kendali manusia (ACL BEA 2025). Sistem rekomendasi berbasis LLM "
        "mencapai performa sebanding dengan model tradisional dalam MOOC recommendation (ArXiv, "
        "2024). Pendekatan goal-based agent menggunakan LLM membantu mahasiswa menavigasi "
        "persyaratan lulus dan prerequisit (Hahsler et al., 2025).")

    add_paragraph(doc,
        "Untuk konteks skill extraction, LLM dapat menangani pola keterampilan kompleks dan "
        "ambiguous melalui few-shot in-context learning, mengatasi keterbatasan reliance pada "
        "training data teranotasi manual (NLP4HR, 2024).")

    # 2.7 Sistem Rekomendasi dalam Pendidikan
    add_heading(doc, "2.7 Sistem Rekomendasi dalam Pendidikan", level=2)

    add_paragraph(doc,
        "Systematic review oleh Zawacki-Richter et al. (2019) terhadap educational recommender "
        "systems mengidentifikasi bahwa pendekatan hybrid paling dominan, dengan collaborative "
        "filtering dan content-based filtering sebagai basis. Machine learning—termasuk SVM, "
        "Naïve Bayes, Decision Trees—dominan dalam implementasi. Namun, gap signifikan ada antara "
        "optimasi metrik relevance dan dampak pedagogis; kurang dari separuh paper mengoptimasi "
        "learning-based metrics (ArXiv, 2024).")

    # 2.8 Kompetensi, Bloom Taxonomy, dan Industry 4.0
    add_heading(doc, "2.8 Kompetensi, Bloom Taxonomy, dan Industry 4.0", level=2)

    add_paragraph(doc,
        "Klasifikasi otomatis learning objectives berdasarkan Bloom's taxonomy telah dieksplorasi "
        "dengan BERT (akurasi hingga 94-95%) dan LLM (EDM 2022, ArXiv 2025). Competency profiles "
        "untuk vocational education perlu selaras dengan Industry 4.0/5.0—lean thinking, data "
        "science, transversal competencies (MDPI, 2023). Systematic review TVET untuk workforce "
        "digital-ready mengidentifikasi manfaat IoT, AI, big data, dan AR/VR untuk pelatihan "
        "(IJMOE, 2024).")

    doc.add_page_break()

    # ========== Kerangka Berpikir ==========
    add_heading(doc, "2.9 Kerangka Berpikir", level=2)

    add_paragraph(doc,
        "Kerangka berpikir penelitian ini digambarkan sebagai alur berikut. Data lowongan kerja "
        "dikumpulkan dari sumber publik. Data tersebut diproses dengan pipeline hybrid NLP: "
        "(1) JobBERT melakukan ekstraksi NER pada level kalimat; (2) LLM melakukan ekstraksi "
        "dengan konteks penuh teks lowongan; (3) hasil digabung dengan fusion logic dan deduplikasi. "
        "Keterampilan yang terekstrak dianalisis tren temporalnya dengan regresi linier dan "
        "kontrol FDR Benjamini-Hochberg. Keterampilan kemudian dipetakan ke domain masa depan "
        "(WEF, O*NET) menggunakan embedding semantik (SBERT) dan trend_score. Prioritas kurikulum "
        "dihitung dari kombinasi demand, trend, dan future_weight. Output sistem—rekomendasi "
        "prioritas keterampilan dan kompetensi—digunakan oleh sekolah untuk mengembangkan KOSP.")

    add_paragraph(doc,
        "Landasan teoretis dari human capital (Becker, 1964; Schultz, 1961) menempatkan sistem "
        "ini sebagai mediator yang memaksimalkan return investasi pendidikan melalui alignment "
        "kurikulum dengan pasar tenaga kerja.")

    doc.add_page_break()

    # ========== BAB III METODE PENELITIAN ==========
    add_heading(doc, "BAB III", level=1)
    add_heading(doc, "METODE PENELITIAN", level=1)

    # 3.1 Pendekatan dan Desain Penelitian
    add_heading(doc, "3.1 Pendekatan dan Desain Penelitian", level=2)

    add_paragraph(doc,
        "Penelitian ini menggunakan Design-Based Research (DBR) sebagai metodologi utama. DBR "
        "adalah pendekatan iteratif yang menjembatani teori dan praktik dalam inovasi pendidikan "
        "(Collins & Brown, 1990; Design-Based Research Collective, 2003). Enam fase DBR—focus, "
        "understand, define, conceive, build, dan test—diterapkan secara rekursif. Peneliti "
        "berkolaborasi dengan praktisi (guru, pengembang kurikulum) sebagai co-researcher.")

    add_paragraph(doc,
        "Pendekatan metode campuran (mixed methods) digunakan: komponen kuantitatif untuk "
        "mengukur efektivitas sistem (precision, AUC, MRR, FDR-controlled trends), dan komponen "
        "kualitatif untuk mengevaluasi persepsi stakeholders melalui expert review dan wawancara.")

    # 3.2 Data dan Sumber Data
    add_heading(doc, "3.2 Data dan Sumber Data", level=2)

    add_paragraph(doc,
        "Data utama berupa lowongan kerja untuk posisi terkait rekayasa perangkat lunak (software "
        "engineer, developer, programmer, dll.) yang dikumpulkan dari sumber publik. Data "
        "dipreprocessing menjadi format job_id, sentence_text, job_date. Gold set untuk evaluasi "
        "dibuat melalui stratified sampling (150 skills, 100 knowledge, 100 domain mapping) "
        "dengan labeling oleh ahli domain.")

    # 3.3 Instrumen dan Teknik Analisis
    add_heading(doc, "3.3 Instrumen dan Teknik Analisis", level=2)

    add_paragraph(doc,
        "Pipeline ekstraksi menggabungkan: (1) JobBERT/SkillSpan-based NER untuk ekstraksi "
        "skill dan knowledge; (2) LLM (misalnya DeepSeek, GPT) untuk ekstraksi dengan konteks "
        "full-text; (3) fusion logic dengan semantic agreement threshold. Analisis tren "
        "menggunakan regresi linier (freq vs month_idx), q-value FDR Benjamini-Hochberg, dan "
        "minimum slope untuk practical significance.")

    add_paragraph(doc,
        "Metrik evaluasi: precision ekstraksi dengan Wilson CI dan binomial test; akurasi "
        "pemetaan domain (Top-1, Top-3, MRR); validasi Bloom dan tipe; AUC-ROC dan Brier score "
        "untuk kalibrasi confidence; Cohen's Kappa/Fleiss' Kappa untuk inter-rater reliability.")

    doc.add_page_break()

    # ========== DAFTAR PUSTAKA ==========
    add_heading(doc, "DAFTAR PUSTAKA", level=1)

    refs = [
        "Becker, G. S. (1964). Human capital: A theoretical and empirical analysis with special reference to education. The University of Chicago Press.",
        "Benjamini, Y., & Hochberg, Y. (1995). Controlling the false discovery rate: A practical and powerful approach to multiple testing. Journal of the Royal Statistical Society B, 57(1), 289-300.",
        "BPS. (2024). Tingkat pengangguran terbuka menurut pendidikan (Agustus 2024). Badan Pusat Statistik. https://www.bps.go.id/",
        "Collins, A., & Brown, J. S. (1990). A new framework for cognition. In L. B. Resnick (Ed.), Knowing, learning, and instruction: Essays in honor of Robert Glaser. Erlbaum.",
        "Design-Based Research Collective. (2003). Design-based research: An emerging paradigm for educational inquiry. Educational Researcher, 32(1), 5-8.",
        "Hahsler, M., et al. (2025). Academic recommender systems using large language models. Research Report.",
        "IJMOE. (2024). Enhancing TVET for a digital-ready workforce: A systematic literature review. International Journal of Modern Education.",
        "Jurnal Online Informatika. (2024). The application of AI technology in vocational high school curriculum design based on individual student skills. JOIN.",
        "Kemendikbudristek. (2022). Peraturan Menteri Pendidikan, Kebudayaan, Riset, dan Teknologi Nomor 56 Tahun 2022 tentang Pedoman Penerapan Kurikulum dalam Rangka Pemulihan Pembelajaran.",
        "MDPI. (2024). Artificial intelligence in curriculum design: A data-driven approach to higher education innovation. MDPI Education Sciences.",
        "Newhouse, D., & Suryadarma, D. (2011). The labor market effects of vocational training in Indonesia. World Bank Policy Research Working Paper.",
        "Permendikbudristek. (2022). Peraturan Menteri Pendidikan, Kebudayaan, Riset, dan Teknologi Nomor 262 Tahun 2022 tentang Perubahan atas Permendikbudristek Nomor 56 Tahun 2022.",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. Proceedings of EMNLP.",
        "Saputri, N. E., & Sulistyawati, A. (2024). Charting vocational education: Impact of agglomeration economies on job-education mismatch in Indonesia. Asia-Pacific Journal of Regional Science.",
        "Schultz, T. W. (1961). Investment in human capital. American Economic Review, 51(1), 1-17.",
        "Senger, E., Zhang, M., van der Goot, R., & Plank, B. (2024). Deep learning-based computational job market analysis: A survey on skill extraction and classification from job postings. Proceedings of NLP4HR 2024.",
        "Springer. (2022). A systematic literature review on educational recommender systems for teaching and learning. Education and Information Technologies.",
        "WJAETS. (2024). Optimizing curriculum design with AI: Aligning educational content with industry demands. World Journal of Advanced Engineering and Technology.",
        "World Economic Forum. (2025). The future of jobs report 2025. WEF. https://www.weforum.org/publications/the-future-of-jobs-report-2025/",
        "Zhang, M., Jensen, K., Sonniks, S., & Plank, B. (2022). SkillSpan: Hard and soft skill extraction from English job postings. Proceedings of NAACL 2022.",
        "Zawacki-Richter, O., Marín, V. I., Bond, M., & Gouverneur, F. (2019). Systematic review of research on artificial intelligence applications in higher education. International Review of Research in Open and Distributed Learning, 20(1), 1-27.",
        "ILO. (2022). Using online vacancy and job applicants' data to study skills dynamics. ILO Working Paper.",
        "ACL BEA. (2025). LLM-assisted, iterative curriculum writing: A human-centered AI approach in Finnish higher education. Proceedings of BEA@ACL.",
        "ArXiv. (2024). From interests to insights: An LLM approach to course recommendations using natural language queries.",
        "ArXiv. (2024). Learning outcomes, assessment, and evaluation in educational recommender systems: A systematic review.",
        "EDM. (2022). Automatic classification of learning objectives based on Bloom's taxonomy. Proceedings of EDM.",
        "Discover AI. (2025). Artificial intelligence-based personalised learning in education: A systematic literature review.",
        "MDPI Applied Sciences. (2023). The role of competence profiles in Industry 5.0-related vocational education and training.",
        "IEEE. (2024). AI-augmented skill development roadmaps for future-ready careers in Education 4.0 and Industry 4.0.",
        "JIEMAR. (2024). Implementation of the Merdeka Curriculum to improve learning outcomes in vocational high schools. Journal of Industrial Engineering & Management Research.",
        "Aktivisme. (2024). Transformasi kurikulum dan standar kompetensi SMK dalam era Merdeka Belajar. Jurnal Ilmu Pendidikan, Politik dan Sosial Indonesia.",
    ]

    for ref in refs:
        add_ref_list_item(doc, ref)

    # Save
    out_path = Path(__file__).resolve().parent.parent / "docs" / "proposal_penelitian.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[INFO] Proposal saved to {out_path}")


if __name__ == "__main__":
    main()
